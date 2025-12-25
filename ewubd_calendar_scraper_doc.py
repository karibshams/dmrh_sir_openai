# ewubd_full_calendar_to_doc.py

import requests
from bs4 import BeautifulSoup
from docx import Document
import os

OUTPUT_FILE = "EWU_Academic_Calendar_FULL.docx"

CALENDAR_LINKS = [

"https://ewubd.edu/academic-calendar-details/spring-2026-schedule-final-exam",
"https://ewubd.edu/academic-calendar-details/spring-2026",
"https://ewubd.edu/academic-calendar-details/spring-2026-2",
"https://ewubd.edu/academic-calendar-details/b-pharm-llb-spring-2026",
"https://ewubd.edu/academic-calendar-details/m-pharm-llm-spring-2026",


"https://ewubd.edu/academic-calendar-details/spring-2025-schedule-final-exam",
"https://ewubd.edu/academic-calendar-details/spring-undergraduate-2025",
"https://ewubd.edu/academic-calendar-details/summer-2025-undergraduate-programs",
"https://ewubd.edu/academic-calendar-details/summer-2025-schedule-final-exam",
"https://ewubd.edu/academic-calendar-details/fall-2025-undergraduate",
"https://ewubd.edu/academic-calendar-details/fall-2025-schedule-final-exam",
"https://ewubd.edu/academic-calendar-details/spring-graduate-2025",
"https://ewubd.edu/academic-calendar-details/summer-2025-graduate",
"https://ewubd.edu/academic-calendar-details/fall-2025",
"https://ewubd.edu/academic-calendar-details/b-pharm-llb-spring-2025",
"https://ewubd.edu/academic-calendar-details/m-pharm-llm-spring-2025",
"https://ewubd.edu/academic-calendar-details/b-pharm-llb-fall-2025",
"https://ewubd.edu/academic-calendar-details/m-pharm-llm-fall-2025",





 
]

def scrape_full_page(url):
    res = requests.get(url, timeout=20)
    res.raise_for_status()
    return BeautifulSoup(res.text, "html.parser")

def append_page_to_doc(soup, doc, url):
    # Page title
    title = soup.find("h2") or soup.find("h1")
    doc.add_heading(title.get_text(strip=True) if title else url, level=2)

    # All visible text blocks (p, h3, h4, strong)
    for tag in soup.find_all(["p", "h3", "h4"]):
        text = tag.get_text(strip=True)
        if text:
            doc.add_paragraph(text)

    # Tables
    for table in soup.find_all("table"):
        rows = table.find_all("tr")
        if not rows:
            continue

        col_count = len(rows[0].find_all(["th", "td"]))
        doc_table = doc.add_table(rows=0, cols=col_count)

        for tr in rows:
            cells = tr.find_all(["th", "td"])
            row_cells = doc_table.add_row().cells
            for i, cell in enumerate(cells):
                row_cells[i].text = cell.get_text(strip=True)

        doc.add_paragraph("")  # spacing

def main():
    if os.path.exists(OUTPUT_FILE):
        doc = Document(OUTPUT_FILE)
    else:
        doc = Document()
        doc.add_heading("East West University – Academic Calendar (Complete)", level=1)

    for link in CALENDAR_LINKS:
        soup = scrape_full_page(link)
        append_page_to_doc(soup, doc, link)

    doc.save(OUTPUT_FILE)
    print(f"Saved all content into {OUTPUT_FILE}")

if __name__ == "__main__":
    main()
