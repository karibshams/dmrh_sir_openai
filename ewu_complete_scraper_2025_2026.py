import requests
from bs4 import BeautifulSoup
from docx import Document
import os

OUTPUT_FILE = "EWU_Complete_Data_2025_2026.docx"


EWU_LINKS = [

    "https://www.ewubd.edu/graduate-programs",
    "https://www.ewubd.edu/undergraduate-programs",

    "https://admission.ewubd.edu/",
    "https://www.ewubd.edu/graduate-programs-tuition-fees",
    "https://www.ewubd.edu/undergraduate-tuition-fees",
    "https://www.ewubd.edu/scholarships-financial-aid",
 
    "https://www.ewubd.edu/board-trustees",
    "https://www.ewubd.edu/vice-chancellor",
    "https://www.ewubd.edu/treasurer",
    "https://www.ewubd.edu/pro-vice-chancellor",
    "https://www.ewubd.edu/registrar",
    "https://www.ewubd.edu/academic-council",
]


def scrape_full_page(url):
    """Fetch and parse a page"""
    print(f"Scraping: {url}")
    headers = {'User-Agent': 'Mozilla/5.0'}
    res = requests.get(url, headers=headers, timeout=20)
    res.raise_for_status()
    return BeautifulSoup(res.text, "html.parser")


def append_page_to_doc(soup, doc, url):
    """Extract all content from page and append to document"""

    title = soup.find("h1") or soup.find("h2") or soup.find("title")
    doc.add_heading(title.get_text(strip=True) if title else url, level=1)

    doc.add_paragraph(f"Source: {url}", style='Intense Quote')
    doc.add_paragraph("") 

    for tag in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "strong"]):
        text = tag.get_text(strip=True)
        if not text or len(text) < 3:
            continue

        if tag.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(tag.name[1])
            doc.add_heading(text, level=min(level + 1, 9))  
  
        elif tag.name == "p":
            if len(text) > 10:  
                doc.add_paragraph(text)
        
        elif tag.name == "strong":
            if len(text) > 5:
                doc.add_paragraph(text, style='Intense Quote')
    
    for ul in soup.find_all(["ul", "ol"]):
        for li in ul.find_all("li", recursive=False):
            text = li.get_text(strip=True)
            if text:
                doc.add_paragraph(text, style='List Bullet')
    
    doc.add_paragraph("") 
    

    for table in soup.find_all("table"):
        rows = table.find_all("tr")
        if not rows:
            continue
        col_count = max(len(row.find_all(["th", "td"])) for row in rows)
        if col_count == 0:
            continue
        
        doc_table = doc.add_table(rows=0, cols=col_count)
        doc_table.style = 'Light Grid Accent 1'
        
        for tr in rows:
            cells = tr.find_all(["th", "td"])
            row_cells = doc_table.add_row().cells
            for i, cell in enumerate(cells):
                if i < col_count:
                    row_cells[i].text = cell.get_text(strip=True)
        
        doc.add_paragraph("") 
    
    doc.add_page_break()


def main():
    """Main scraping function"""
    print("="*70)
    print("EWU Complete Data Scraper 2025-2026")
    print("="*70)
    
    if os.path.exists(OUTPUT_FILE):
        print(f"\nLoading existing document: {OUTPUT_FILE}")
        doc = Document(OUTPUT_FILE)
    else:
        doc = Document()
        doc.add_heading("East West University – Complete Data 2025-2026", level=0)
        doc.add_paragraph("Academic Year: 2025-2026")
        doc.add_paragraph("")
    
    for i, link in enumerate(EWU_LINKS, 1):
        print(f"\n[{i}/{len(EWU_LINKS)}] Processing: {link}")
        
        try:
            soup = scrape_full_page(link)
            append_page_to_doc(soup, doc, link)
            print(f"  ✓ Successfully added to document")
            
        except Exception as e:
            print(f"  ✗ Error: {e}")
            continue
    
    print(f"\n{'='*70}")
    print(f"Saving document to {OUTPUT_FILE}...")
    doc.save(OUTPUT_FILE)
    
    print(f"✓ Successfully saved!")
    print(f"{'='*70}")
    print(f"\n ALL DATA SAVED TO {OUTPUT_FILE}")
    print(f"{'='*70}\n")


if __name__ == "__main__":
    main()